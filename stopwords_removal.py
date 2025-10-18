import streamlit as st
import os
from docx import Document

# Load stopwords once
def load_stopwords(path="sw.txt"):
    with open(path, "r", encoding="utf-8") as f:
        stopwords = set(line.strip() for line in f if line.strip())
    return stopwords

# Remove stopwords but preserve paragraphs
def remove_stopwords(text, stopwords):
    lines = text.splitlines()
    filtered_lines = []
    for line in lines:
        words = line.split()
        filtered = [w for w in words if w not in stopwords]
        filtered_lines.append(" ".join(filtered))
    return "\n".join(filtered_lines)

# Streamlit UI
st.sidebar.image("images/peacock-3.png", width=200)
st.sidebar.markdown("<h3 style='text-align: center;'>NLP Tool</h3>", unsafe_allow_html=True)

# Options
select = ['remove-stopwords']
option = st.sidebar.selectbox('Choose an option', select)

if option == "remove-stopwords":

    uploaded_file = st.file_uploader("📂 Choose a .docx file", type=["docx"])

    if uploaded_file is not None:
        try:
            # Get uploaded file name (without extension)
            base_filename = os.path.splitext(uploaded_file.name)[0]

            # Load the Word document
            doc = Document(uploaded_file)
            full_text = "\n".join([para.text for para in doc.paragraphs])

            # Load stopwords
            stopwords = load_stopwords("sw.txt")

            # Remove stopwords
            result_text = remove_stopwords(full_text, stopwords)

            # Save processed Word document
            os.makedirs("output", exist_ok=True)
            output_filename = f"{base_filename}_output.docx"
            output_path = os.path.join("output", output_filename)

            new_doc = Document()
            for line in result_text.splitlines():
                new_doc.add_paragraph(line)
            new_doc.save(output_path)

            # Provide download button
            with open(output_path, "rb") as f:
                st.download_button(
                    label=f"📄 Download {output_filename}",
                    data=f,
                    file_name=output_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            # Show preview
            st.subheader("📝 Output Preview:")
            st.text(result_text)

        except Exception as e:
            st.error(f"❌ Error processing file: {e}")